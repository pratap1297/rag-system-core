"""
Document Validator for RAG System
Quality assurance and validation for document processing
"""

import os
import mimetypes
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime

from core.logging_system import get_logger
from core.exceptions import ValidationError
from core.error_handler import with_error_handling

@dataclass
class ValidationResult:
    """Document validation result"""
    is_valid: bool
    file_path: str
    file_size: int
    file_type: str
    mime_type: str
    issues: List[str]
    warnings: List[str]
    validation_time: float
    
    @property
    def has_issues(self) -> bool:
        """Check if document has validation issues"""
        return len(self.issues) > 0
    
    @property
    def has_warnings(self) -> bool:
        """Check if document has warnings"""
        return len(self.warnings) > 0

class DocumentValidator:
    """Validates documents for processing"""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        self.config = config or {}
        self.logger = get_logger("document_validator")
        
        # Validation settings
        self.max_file_size = self.config.get('max_file_size', 100 * 1024 * 1024)  # 100MB
        self.min_file_size = self.config.get('min_file_size', 1)  # 1 byte
        self.supported_formats = self.config.get('supported_formats', ['.pdf', '.docx', '.doc', '.txt'])
        self.blocked_formats = self.config.get('blocked_formats', ['.exe', '.bat', '.sh'])
        
        self.logger.info("Document validator initialized")
    
    @with_error_handling("document_validator", "validate_document")
    def validate_document(self, file_path: str) -> ValidationResult:
        """Validate a single document"""
        
        start_time = datetime.now()
        
        try:
            file_path = Path(file_path)
            
            issues = []
            warnings = []
            
            # Basic file existence and accessibility
            if not file_path.exists():
                issues.append("File does not exist")
                
                return ValidationResult(
                    is_valid=False,
                    file_path=str(file_path),
                    file_size=0,
                    file_type="",
                    mime_type="",
                    issues=issues,
                    warnings=warnings,
                    validation_time=0.0
                )
            
            if not file_path.is_file():
                issues.append("Path is not a file")
            
            # File permissions
            if not os.access(file_path, os.R_OK):
                issues.append("File is not readable")
            
            # File size validation
            file_size = file_path.stat().st_size
            
            if file_size < self.min_file_size:
                issues.append(f"File too small: {file_size} bytes (minimum: {self.min_file_size})")
            
            if file_size > self.max_file_size:
                issues.append(f"File too large: {file_size} bytes (maximum: {self.max_file_size})")
            
            # File type validation
            file_extension = file_path.suffix.lower()
            mime_type, _ = mimetypes.guess_type(str(file_path))
            
            if file_extension in self.blocked_formats:
                issues.append(f"Blocked file format: {file_extension}")
            
            if file_extension not in self.supported_formats:
                warnings.append(f"Unsupported file format: {file_extension}")
            
            # Content validation
            content_issues, content_warnings = self._validate_content(file_path)
            issues.extend(content_issues)
            warnings.extend(content_warnings)
            
            # Security validation
            security_issues = self._validate_security(file_path)
            issues.extend(security_issues)
            
            # Calculate validation time
            validation_time = (datetime.now() - start_time).total_seconds()
            
            # Create result
            result = ValidationResult(
                is_valid=len(issues) == 0,
                file_path=str(file_path),
                file_size=file_size,
                file_type=file_extension,
                mime_type=mime_type or "unknown",
                issues=issues,
                warnings=warnings,
                validation_time=validation_time
            )
            
            self.logger.info(f"Validation completed for {file_path.name}: "
                           f"valid={result.is_valid}, issues={len(issues)}, warnings={len(warnings)}")
            
            return result
            
        except Exception as e:
            validation_time = (datetime.now() - start_time).total_seconds()
            
            return ValidationResult(
                is_valid=False,
                file_path=str(file_path),
                file_size=0,
                file_type="",
                mime_type="",
                issues=[f"Validation error: {str(e)}"],
                warnings=[],
                validation_time=validation_time
            )
    
    def validate_batch(self, file_paths: List[str]) -> List[ValidationResult]:
        """Validate multiple documents"""
        
        self.logger.info(f"Starting batch validation of {len(file_paths)} files")
        
        results = []
        valid_count = 0
        
        for i, file_path in enumerate(file_paths):
            try:
                result = self.validate_document(file_path)
                results.append(result)
                
                if result.is_valid:
                    valid_count += 1
                
                self.logger.debug(f"Validated {i+1}/{len(file_paths)}: {Path(file_path).name}")
                
            except Exception as e:
                self.logger.error(f"Failed to validate {file_path}: {e}")
                
                # Create failed validation result
                failed_result = ValidationResult(
                    is_valid=False,
                    file_path=file_path,
                    file_size=0,
                    file_type="",
                    mime_type="",
                    issues=[f"Validation failed: {str(e)}"],
                    warnings=[],
                    validation_time=0.0
                )
                results.append(failed_result)
        
        self.logger.info(f"Batch validation completed: {valid_count}/{len(file_paths)} valid")
        return results
    
    def _validate_content(self, file_path: Path) -> Tuple[List[str], List[str]]:
        """Validate file content"""
        
        issues = []
        warnings = []
        
        try:
            file_extension = file_path.suffix.lower()
            
            # PDF-specific validation
            if file_extension == '.pdf':
                pdf_issues, pdf_warnings = self._validate_pdf(file_path)
                issues.extend(pdf_issues)
                warnings.extend(pdf_warnings)
            
            # Text file validation
            elif file_extension == '.txt':
                txt_issues, txt_warnings = self._validate_text_file(file_path)
                issues.extend(txt_issues)
                warnings.extend(txt_warnings)
            
            # Word document validation
            elif file_extension in ['.docx', '.doc']:
                word_issues, word_warnings = self._validate_word_document(file_path)
                issues.extend(word_issues)
                warnings.extend(word_warnings)
            
        except Exception as e:
            issues.append(f"Content validation error: {str(e)}")
        
        return issues, warnings
    
    def _validate_pdf(self, file_path: Path) -> Tuple[List[str], List[str]]:
        """Validate PDF file"""
        
        issues = []
        warnings = []
        
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                try:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    # Check if PDF is encrypted
                    if pdf_reader.is_encrypted:
                        issues.append("PDF is encrypted/password protected")
                    
                    # Check page count
                    page_count = len(pdf_reader.pages)
                    
                    if page_count == 0:
                        issues.append("PDF has no pages")
                    elif page_count > 1000:
                        warnings.append(f"PDF has many pages ({page_count}), processing may be slow")
                    
                    # Try to extract text from first page
                    if page_count > 0 and not pdf_reader.is_encrypted:
                        try:
                            first_page = pdf_reader.pages[0]
                            text = first_page.extract_text()
                            
                            if not text.strip():
                                warnings.append("PDF appears to contain no extractable text (may be image-based)")
                        
                        except Exception:
                            warnings.append("Could not extract text from PDF")
                
                except PyPDF2.errors.PdfReadError as e:
                    issues.append(f"PDF read error: {str(e)}")
                except Exception as e:
                    issues.append(f"PDF validation error: {str(e)}")
        
        except ImportError:
            warnings.append("PyPDF2 not available for PDF validation")
        except Exception as e:
            issues.append(f"PDF validation failed: {str(e)}")
        
        return issues, warnings
    
    def _validate_text_file(self, file_path: Path) -> Tuple[List[str], List[str]]:
        """Validate text file"""
        
        issues = []
        warnings = []
        
        try:
            # Try to read file with different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            readable = False
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        # Try to read first few lines
                        for _ in range(10):
                            line = file.readline()
                            if not line:
                                break
                    
                    readable = True
                    break
                
                except UnicodeDecodeError:
                    continue
                except Exception:
                    break
            
            if not readable:
                issues.append("Text file cannot be decoded with common encodings")
            
        except Exception as e:
            issues.append(f"Text file validation error: {str(e)}")
        
        return issues, warnings
    
    def _validate_word_document(self, file_path: Path) -> Tuple[List[str], List[str]]:
        """Validate Word document"""
        
        issues = []
        warnings = []
        
        try:
            import docx
            
            try:
                doc = docx.Document(file_path)
                
                # Check if document has content
                paragraph_count = len(doc.paragraphs)
                
                if paragraph_count == 0:
                    warnings.append("Word document appears to have no paragraphs")
                
                # Check for text content
                has_text = False
                for paragraph in doc.paragraphs[:10]:  # Check first 10 paragraphs
                    if paragraph.text.strip():
                        has_text = True
                        break
                
                if not has_text:
                    warnings.append("Word document appears to have no text content")
            
            except docx.opc.exceptions.PackageNotFoundError:
                issues.append("Invalid Word document format")
            except Exception as e:
                issues.append(f"Word document validation error: {str(e)}")
        
        except ImportError:
            warnings.append("python-docx not available for Word document validation")
        except Exception as e:
            issues.append(f"Word document validation failed: {str(e)}")
        
        return issues, warnings
    
    def _validate_security(self, file_path: Path) -> List[str]:
        """Validate file security"""
        
        issues = []
        
        try:
            # Check file name for suspicious patterns
            file_name = file_path.name.lower()
            
            suspicious_patterns = [
                'script', 'macro', 'vba', 'javascript', 'powershell',
                'cmd', 'bat', 'exe', 'dll', 'scr'
            ]
            
            for pattern in suspicious_patterns:
                if pattern in file_name:
                    issues.append(f"Suspicious file name pattern: {pattern}")
                    break
            
            # Check for double extensions
            if file_name.count('.') > 1:
                parts = file_name.split('.')
                if len(parts) > 2 and parts[-2] in ['exe', 'bat', 'cmd', 'scr']:
                    issues.append("Suspicious double file extension")
            
        except Exception as e:
            self.logger.warning(f"Security validation error: {e}")
        
        return issues
    
    def get_validation_summary(self, results: List[ValidationResult]) -> Dict[str, Any]:
        """Get validation summary statistics"""
        
        total_files = len(results)
        valid_files = sum(1 for r in results if r.is_valid)
        invalid_files = total_files - valid_files
        
        total_issues = sum(len(r.issues) for r in results)
        total_warnings = sum(len(r.warnings) for r in results)
        
        # File type distribution
        file_types = {}
        for result in results:
            file_type = result.file_type or 'unknown'
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        # Common issues
        issue_counts = {}
        for result in results:
            for issue in result.issues:
                issue_counts[issue] = issue_counts.get(issue, 0) + 1
        
        return {
            'total_files': total_files,
            'valid_files': valid_files,
            'invalid_files': invalid_files,
            'validation_rate': (valid_files / total_files * 100) if total_files > 0 else 0,
            'total_issues': total_issues,
            'total_warnings': total_warnings,
            'file_types': file_types,
            'common_issues': dict(sorted(issue_counts.items(), key=lambda x: x[1], reverse=True)[:10])
        } 